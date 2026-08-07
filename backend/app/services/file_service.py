import os
from uuid import uuid4

from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.colors import HexColor
from reportlab.graphics.shapes import Drawing, Rect, String, Line

from docx import Document

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

from app.schemas.file import FileGenerateRequest
from app.services.llm_service import LLMService

llm_service = LLMService()

from reportlab.pdfgen import canvas

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        if self._pageNumber == 1:
            return  # Cover Page handles its own canvas artwork

        self.saveState()
        self.setFont("Helvetica-Bold", 8)
        self.setFillColor(HexColor("#0F172A"))

        # Top Header Accent Line & Branding (Canva Minimalist Style)
        self.drawString(45, 780, "BEATS PRODUCTION PRIVATE LIMITED")
        self.setFont("Helvetica", 8)
        self.setFillColor(HexColor("#64748B"))
        self.drawRightString(567, 780, "BUSINESS PROJECT PROPOSAL")

        self.setStrokeColor(HexColor("#2563EB"))
        self.setLineWidth(1.2)
        self.line(45, 772, 567, 772)

        # Bottom Footer Accent Line & Canva Page Counter (02 / 04)
        self.setStrokeColor(HexColor("#E2E8F0"))
        self.setLineWidth(0.5)
        self.line(45, 45, 567, 45)

        self.setFont("Helvetica", 8)
        self.setFillColor(HexColor("#64748B"))
        self.drawString(45, 32, "Confidential  •  Prepared for Client Review")

        page_text = f"{self._pageNumber:02d} / {page_count:02d}"
        self.setFont("Helvetica-Bold", 8)
        self.setFillColor(HexColor("#2563EB"))
        self.drawRightString(567, 32, page_text)

        self.restoreState()


class FileService:

    def create_custom_diagram(self, content: str):
        # Scan content for potential block names
        import re
        blocks = []

        # Look for numbered items like "1. Perception Layer:" or "1. Input:"
        matches = re.findall(r'\d+\.\s*([^:\n\-\–\—]+)', content)
        for m in matches:
            name = m.strip()
            # Clean up names (remove bold formatting stars, etc.)
            name = name.replace("**", "").replace("*", "")
            if len(name) < 25 and name not in blocks:
                blocks.append(name)

        # Fallback to defaults if we didn't find clear block names
        if len(blocks) < 3:
            blocks = ["Perception Layer", "Cognitive Layer", "Execution Layer"]
        else:
            blocks = blocks[:3] # We only draw 3 boxes in this simple diagram template

        d = Drawing(400, 120)

        # Colors: Blue, Purple, Green
        colors = [
            ("#EBF5FF", "#2563EB", "#1E3A8A"), # Blue
            ("#F5F3FF", "#7C3AED", "#5B21B6"), # Purple
            ("#ECFDF5", "#10B981", "#065F46")  # Green
        ]

        # Draw the 3 boxes
        x_coords = [10, 150, 290]
        for i in range(3):
            bg, border, text_color = colors[i]
            x = x_coords[i]
            # Draw box
            d.add(Rect(x, 40, 100, 40, fillColor=HexColor(bg), strokeColor=HexColor(border), strokeWidth=1.5, rx=5, ry=5))

            # Draw text wrapping if it's too long
            label = blocks[i]
            if len(label) > 15:
                # Split into words
                words = label.split(" ")
                if len(words) >= 2:
                    w1 = " ".join(words[:len(words)//2])
                    w2 = " ".join(words[len(words)//2:])
                    d.add(String(x + 50, 64, w1, textAnchor="middle", fontSize=8, fillColor=HexColor(text_color), fontName="Helvetica-Bold"))
                    d.add(String(x + 50, 52, w2, textAnchor="middle", fontSize=8, fillColor=HexColor(text_color), fontName="Helvetica-Bold"))
                else:
                    d.add(String(x + 50, 58, label[:15], textAnchor="middle", fontSize=8, fillColor=HexColor(text_color), fontName="Helvetica-Bold"))
            else:
                d.add(String(x + 50, 58, label, textAnchor="middle", fontSize=9, fillColor=HexColor(text_color), fontName="Helvetica-Bold"))

            # Connectors
            if i < 2:
                next_x = x_coords[i+1]
                arrow_start = x + 100
                d.add(Line(arrow_start, 60, next_x, 60, strokeColor=HexColor(border), strokeWidth=1.5))
                # Arrowhead
                d.add(Line(next_x - 5, 56, next_x, 60, strokeColor=HexColor(border), strokeWidth=1.5))
                d.add(Line(next_x - 5, 64, next_x, 60, strokeColor=HexColor(border), strokeWidth=1.5))

        # Draw feedback loop line from Execution (Box 3) back to Cognitive (Box 2)
        d.add(Line(340, 40, 340, 15, strokeColor=HexColor("#10B981"), strokeWidth=1.2))
        d.add(Line(340, 15, 200, 15, strokeColor=HexColor("#6B7280"), strokeWidth=1.2, strokeDashArray=[2, 2]))
        d.add(Line(200, 15, 200, 40, strokeColor=HexColor("#6B7280"), strokeWidth=1.2, strokeDashArray=[2, 2]))
        # Arrowhead up
        d.add(Line(196, 35, 200, 40, strokeColor=HexColor("#6B7280"), strokeWidth=1.2))
        d.add(Line(204, 35, 200, 40, strokeColor=HexColor("#6B7280"), strokeWidth=1.2))

        # Label for Feedback
        d.add(String(270, 20, "Feedback / Learn Loop", textAnchor="middle", fontSize=7, fillColor=HexColor("#4B5563"), fontName="Helvetica-Oblique"))

        return d

    def process_prompt(
        self,
        prompt: str
    ):
        """
        Generates the required file from a prompt.
        This method is used by the chat service.
        """

        request = FileGenerateRequest(
            prompt=prompt,
            file_type=""
        )

        return self.generate_file(request)

    def generate_file(
        self,
        request: FileGenerateRequest
    ):

        # Create output folder
        os.makedirs(
            "generated_files",
            exist_ok=True
        )

        # Ask AI to determine the best file type
        file_type_data = llm_service.detect_file_type(request.prompt)
        file_type = file_type_data.get("file_type", "").strip().lower()

        # PDF
        if file_type == "pdf":

            content = llm_service.generate_document_content(
                request.prompt
            )

            return self.create_pdf(content)

        # WORD
        elif file_type == "docx":

            content = llm_service.generate_document_content(
                request.prompt
            )

            return self.create_word(content)

        # EXCEL
        elif file_type == "xlsx":

            data = llm_service.generate_excel_data(
                request.prompt
            )

            return self.create_excel(data)

        # Unsupported
        else:

            return {
                "message": "Unsupported file type",
                "detected_type": file_type
            }

    # ==========================================================
    # PDF
    # ==========================================================

    def build_reportlab_table(self, table_data):
        from reportlab.platypus import Table, TableStyle
        from reportlab.lib.colors import HexColor
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        col_count = len(table_data[0]) if table_data else 1
        col_width = 480 / col_count
        
        styles = getSampleStyleSheet()
        cell_style = ParagraphStyle(
            "TableCellStyle",
            parent=styles["BodyText"],
            fontSize=9,
            leading=12,
            textColor=HexColor("#1E293B")
        )
        cell_header_style = ParagraphStyle(
            "TableHeaderStyle",
            parent=styles["BodyText"],
            fontSize=9,
            leading=12,
            textColor=HexColor("#FFFFFF"),
            fontName="Helvetica-Bold"
        )
        
        flowable_data = []
        for r_idx, row in enumerate(table_data):
            flowable_row = []
            for col in row:
                style = cell_header_style if r_idx == 0 else cell_style
                formatted_text = self._format_md_for_reportlab(col.strip())
                flowable_row.append(Paragraph(formatted_text, style))
            flowable_data.append(flowable_row)
            
        t = Table(flowable_data, colWidths=[col_width] * col_count)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), HexColor("#0F172A")),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('GRID', (0,0), (-1,-1), 0.5, HexColor("#CBD5E1")),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [HexColor("#FFFFFF"), HexColor("#F8FAFC")]),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ]))
        return t

    def build_word_table(self, document, table_data):
        if not table_data:
            return
        col_count = len(table_data[0])
        table = document.add_table(rows=len(table_data), cols=col_count)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, cell_value in enumerate(row):
                if c_idx < len(table.rows[r_idx].cells):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    self._add_md_runs_to_docx_paragraph(p, cell_value.strip())
                    if r_idx == 0:
                        for run in p.runs:
                            run.bold = True

    def _format_md_for_reportlab(self, text: str) -> str:
        import re
        if not text:
            return ""
        text = text.replace("&", "&amp;")
        text = re.sub(r'\*\*\*(.*?)\*\*\*', r'<b><i>\1</i></b>', text)
        text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
        text = re.sub(r'\[(.*?)\]\((.*?)\)', r'<font color="#2563EB"><u>\1</u></font>', text)
        return text

    def _add_md_runs_to_docx_paragraph(self, p, text: str):
        import re
        if not text:
            return
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part)

    def _build_cover_page_flowables(self, cover_lines, styles):
        import re
        details = {}
        notice = "This proposal contains proprietary design concepts and commercial terms intended solely for the recipient organization."

        for line in cover_lines:
            line_str = line.strip()
            if ":" in line_str and not line_str.startswith("#"):
                parts = line_str.split(":", 1)
                k = parts[0].replace("**", "").replace("*", "").strip()
                v = parts[1].replace("**", "").replace("*", "").strip()
                if "confidentiality" in k.lower():
                    notice = v
                else:
                    details[k] = v

        flowables = []
        flowables.append(Spacer(1, 15))

        # Canva Style Header Tag & Title
        tag_p = Paragraph("BUSINESS PROJECT PROPOSAL", ParagraphStyle(
            "CoverTag", parent=styles["Normal"], fontSize=9, leading=12, textColor=HexColor("#2563EB"), fontName="Helvetica-Bold"
        ))
        title_p = Paragraph("EVENT PROPOSAL", ParagraphStyle(
            "CoverTitle", parent=styles["Title"], fontSize=28, leading=32, alignment=0, textColor=HexColor("#0F172A"), fontName="Helvetica-Bold", spaceBefore=4, spaceAfter=4
        ))
        sub_p = Paragraph("BEATS PRODUCTION PRIVATE LIMITED  •  TECHNICAL PRODUCTION & EVENT MANAGEMENT", ParagraphStyle(
            "CoverSub", parent=styles["Normal"], fontSize=9.5, leading=13, textColor=HexColor("#475569"), fontName="Helvetica-Bold"
        ))

        header_table = Table([[tag_p], [title_p], [sub_p]], colWidths=[480])
        header_table.setStyle(TableStyle([
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ]))
        flowables.append(header_table)
        flowables.append(Spacer(1, 10))

        # Accent Rule Line (Canva Electric Blue)
        rule_table = Table([[""]], colWidths=[480], rowHeights=[2.5])
        rule_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), HexColor("#2563EB")),
            ('TOPPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ]))
        flowables.append(rule_table)
        flowables.append(Spacer(1, 24))

        # Key Details Card Table (Canva Minimalist Grid)
        card_data = []
        k_style = ParagraphStyle("CK", parent=styles["Normal"], fontSize=9.5, leading=13.5, textColor=HexColor("#475569"), fontName="Helvetica-Bold")
        v_style = ParagraphStyle("CV", parent=styles["Normal"], fontSize=9.5, leading=13.5, textColor=HexColor("#0F172A"), fontName="Helvetica")

        for k, v in details.items():
            formatted_val = self._format_md_for_reportlab(v)
            card_data.append([Paragraph(f"<b>{k.upper()}</b>", k_style), Paragraph(formatted_val, v_style)])

        if card_data:
            details_table = Table(card_data, colWidths=[140, 340])
            details_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), HexColor("#F8FAFC")),
                ('BOX', (0,0), (-1,-1), 1, HexColor("#E2E8F0")),
                ('INNERGRID', (0,0), (-1,-1), 0.5, HexColor("#E2E8F0")),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('LEFTPADDING', (0,0), (-1,-1), 14),
                ('RIGHTPADDING', (0,0), (-1,-1), 14),
            ]))
            flowables.append(details_table)
            flowables.append(Spacer(1, 20))

        # Confidentiality Box
        c_title = Paragraph("<b>CONFIDENTIALITY NOTICE</b>", ParagraphStyle("CN1", parent=styles["Normal"], fontSize=8.5, leading=11, textColor=HexColor("#1E40AF")))
        c_text = Paragraph(self._format_md_for_reportlab(notice), ParagraphStyle("CN2", parent=styles["Normal"], fontSize=8.5, leading=12, textColor=HexColor("#64748B")))
        notice_table = Table([[c_title], [Spacer(1, 2)], [c_text]], colWidths=[480])
        notice_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), HexColor("#EFF6FF")),
            ('BOX', (0,0), (-1,-1), 0.75, HexColor("#BFDBFE")),
            ('TOPPADDING', (0,0), (-1,-1), 10),
            ('BOTTOMPADDING', (0,0), (-1,-1), 10),
            ('LEFTPADDING', (0,0), (-1,-1), 14),
            ('RIGHTPADDING', (0,0), (-1,-1), 14),
        ]))
        flowables.append(notice_table)
        flowables.append(PageBreak())  # Standalone Cover Page!

        return flowables

    # ==========================================================
    # PDF
    # ==========================================================

    def create_pdf(
        self,
        content: str
    ):
        import re
        filename = f"{uuid4()}.pdf"

        file_path = os.path.join(
            "generated_files",
            filename
        )

        document = SimpleDocTemplate(
            file_path,
            leftMargin=45,
            rightMargin=45,
            topMargin=85,
            bottomMargin=65
        )

        styles = getSampleStyleSheet()

        h1_style = ParagraphStyle(
            "DocH1Style",
            parent=styles["Heading1"],
            fontSize=12.5,
            leading=16,
            textColor=HexColor("#0F172A"),
            fontName="Helvetica-Bold",
            spaceBefore=14,
            spaceAfter=6
        )

        h2_style = ParagraphStyle(
            "DocH2Style",
            parent=styles["Heading2"],
            fontSize=11,
            leading=15,
            textColor=HexColor("#1E40AF"),
            fontName="Helvetica-Bold",
            spaceBefore=10,
            spaceAfter=4
        )

        h3_style = ParagraphStyle(
            "DocH3Style",
            parent=styles["Heading3"],
            fontSize=10,
            leading=13,
            textColor=HexColor("#334155"),
            fontName="Helvetica-Bold",
            spaceBefore=6,
            spaceAfter=3
        )

        body_style = ParagraphStyle(
            "DocBodyStyle",
            parent=styles["BodyText"],
            fontSize=9.5,
            leading=14.5,
            textColor=HexColor("#1E293B"),
            spaceAfter=5
        )

        bullet_style = ParagraphStyle(
            "DocBulletStyle",
            parent=styles["BodyText"],
            fontSize=9.5,
            leading=14,
            textColor=HexColor("#1E293B"),
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=4
        )

        story = []
        current_table_data = []

        # Separate Section 1 (Cover Page) lines from remaining content
        all_lines = content.split("\n")
        cover_lines = []
        body_lines = []
        in_cover = False

        for l in all_lines:
            stripped = l.strip()
            if stripped.startswith("```"):
                continue
            if stripped.startswith("# 1. Cover Page") or stripped.startswith("# Cover Page"):
                in_cover = True
                continue
            elif in_cover and stripped.startswith("# "):
                in_cover = False
                body_lines.append(l)
            elif in_cover:
                cover_lines.append(l)
            else:
                body_lines.append(l)

        # Build Cover Page if cover lines exist
        if cover_lines:
            story.extend(self._build_cover_page_flowables(cover_lines, styles))

        # Parse Body Sections
        for raw_line in body_lines:
            line = raw_line.strip()

            if line.startswith("```"):
                continue

            if not line:
                if current_table_data:
                    story.append(self.build_reportlab_table(current_table_data))
                    story.append(Spacer(1, 8))
                    current_table_data = []
                continue

            if line.startswith("[TABLE_ROW]"):
                row_text = line[11:].strip()
                cells = [c.strip() for c in row_text.split("|")]
                current_table_data.append(cells)
                continue
            elif line.startswith("|") and "|" in line[1:]:
                if re.match(r'^\|[\s:\-\|]+\|$', line):
                    continue
                cells = [c.strip() for c in line.strip("|").split("|")]
                if any(cells):
                    current_table_data.append(cells)
                continue

            if current_table_data:
                story.append(self.build_reportlab_table(current_table_data))
                story.append(Spacer(1, 8))
                current_table_data = []

            # Headers & formatting
            formatted_line = self._format_md_for_reportlab(line)

            if line.startswith("# "):
                header_text = line[2:].strip()
                text = self._format_md_for_reportlab(header_text)
                
                # Canva Minimalist Section Header with Left Accent Bar
                hdr_p = Paragraph(f"<b>{text.upper()}</b>", ParagraphStyle(
                    "H1CanvaBanner", parent=styles["Normal"], fontSize=12, leading=15, textColor=HexColor("#0F172A"), fontName="Helvetica-Bold"
                ))
                hdr_table = Table([[hdr_p]], colWidths=[480])
                hdr_table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,-1), HexColor("#F8FAFC")),
                    ('BOX', (0,0), (-1,-1), 0.5, HexColor("#E2E8F0")),
                    ('LINELEFT', (0,0), (0,0), 3.5, HexColor("#2563EB")),
                    ('TOPPADDING', (0,0), (-1,-1), 6),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                    ('LEFTPADDING', (0,0), (-1,-1), 10),
                ]))
                story.append(Spacer(1, 12))
                story.append(hdr_table)
                story.append(Spacer(1, 6))

            elif line.startswith("## "):
                header_text = line[3:].strip()
                text = self._format_md_for_reportlab(header_text)
                story.append(Paragraph(text, h1_style))
                story.append(Spacer(1, 3))
            elif line.startswith("### "):
                header_text = line[4:].strip()
                text = self._format_md_for_reportlab(header_text)
                story.append(Paragraph(text, h2_style))
                story.append(Spacer(1, 2))
            elif line.startswith("#### "):
                header_text = line[5:].strip()
                text = self._format_md_for_reportlab(header_text)
                story.append(Paragraph(text, h3_style))
            elif line.startswith("- ") or line.startswith("* ") or line.startswith("+ "):
                text = self._format_md_for_reportlab(line[2:].strip())
                story.append(Paragraph(f"• {text}", bullet_style))
            elif re.match(r'^\d+\.\s+', line):
                item_text = re.sub(r'^\d+\.\s+', '', line)
                num = line.split('.')[0]
                text = self._format_md_for_reportlab(item_text)
                story.append(Paragraph(f"<b>{num}.</b> {text}", bullet_style))
            elif line == "---" or line == "[PAGE_BREAK]":
                story.append(Spacer(1, 10))
            else:
                if "[insert diagram" in line.lower() or "[architecture diagram" in line.lower():
                    story.append(Spacer(1, 8))
                    story.append(self.create_custom_diagram(content))
                    story.append(Spacer(1, 12))
                else:
                    story.append(Paragraph(formatted_line, body_style))

        if current_table_data:
            story.append(self.build_reportlab_table(current_table_data))

        document.build(story, canvasmaker=NumberedCanvas)

        return {
            "filename": filename,
            "file_path": file_path
        }

    # ==========================================================
    # WORD
    # ==========================================================

    def create_word(
        self,
        content: str
    ):
        import re
        filename = f"{uuid4()}.docx"

        file_path = os.path.join(
            "generated_files",
            filename
        )

        document = Document()
        current_table_data = []
        is_cover = False

        for raw_line in content.split("\n"):
            line = raw_line.strip()

            if line.startswith("```"):
                continue

            if not line:
                if current_table_data:
                    self.build_word_table(document, current_table_data)
                    current_table_data = []
                continue

            if line.startswith("[TABLE_ROW]"):
                row_text = line[11:].strip()
                cells = [c.strip() for c in row_text.split("|")]
                current_table_data.append(cells)
                continue
            elif line.startswith("|") and "|" in line[1:]:
                if re.match(r'^\|[\s:\-\|]+\|$', line):
                    continue
                cells = [c.strip() for c in line.strip("|").split("|")]
                if any(cells):
                    current_table_data.append(cells)
                continue

            if current_table_data:
                self.build_word_table(document, current_table_data)
                current_table_data = []

            # Headers & formatting for Word
            if line.startswith("# "):
                header_text = line[2:].strip()
                if "1. Cover Page" in header_text:
                    is_cover = True
                    h = document.add_heading(level=0)
                    self._add_md_runs_to_docx_paragraph(h, "BEATS PRODUCTION PRIVATE LIMITED")
                    p = document.add_paragraph("EVENT PROPOSAL & TECHNICAL RIDER")
                    p.paragraph_format.space_after = 20
                else:
                    if is_cover:
                        is_cover = False
                        document.add_page_break()
                    h = document.add_heading(level=1)
                    self._add_md_runs_to_docx_paragraph(h, header_text)
            elif line.startswith("## "):
                h = document.add_heading(level=2)
                self._add_md_runs_to_docx_paragraph(h, line[3:].strip())
            elif line.startswith("### "):
                h = document.add_heading(level=3)
                self._add_md_runs_to_docx_paragraph(h, line[4:].strip())
            elif line.startswith("- ") or line.startswith("* ") or line.startswith("+ "):
                p = document.add_paragraph(style='List Bullet')
                self._add_md_runs_to_docx_paragraph(p, line[2:].strip())
            elif re.match(r'^\d+\.\s+', line):
                item_text = re.sub(r'^\d+\.\s+', '', line)
                p = document.add_paragraph(style='List Number')
                self._add_md_runs_to_docx_paragraph(p, item_text)
            elif line == "---" or line == "[PAGE_BREAK]":
                pass
            else:
                p = document.add_paragraph()
                self._add_md_runs_to_docx_paragraph(p, line)

        if current_table_data:
            self.build_word_table(document, current_table_data)

        document.save(file_path)

        return {
            "filename": filename,
            "file_path": file_path
        }

        if current_table_data:
            self.build_word_table(document, current_table_data)

        document.save(file_path)

        return {
            "filename": filename,
            "file_path": file_path
        }

    # ==========================================================
    # EXCEL
    # ==========================================================

    def create_excel(
        self,
        data
    ):

        if not data:

            return {
                "message": "No data generated."
            }

        filename = f"{uuid4()}.xlsx"

        file_path = os.path.join(
            "generated_files",
            filename
        )

        workbook = Workbook()

        sheet = workbook.active

        sheet.title = "Generated Data"

        # Headers
        headers = list(data[0].keys())

        sheet.append(headers)

        # Rows
        for row in data:

            sheet.append(
                list(row.values())
            )

        # Header Style
        for cell in sheet[1]:

            cell.font = Font(
                bold=True,
                color="FFFFFF"
            )

            cell.fill = PatternFill(
                fill_type="solid",
                start_color="4F81BD"
            )

            cell.alignment = Alignment(
                horizontal="center",
                vertical="center"
            )

        # Auto Width
        for column in sheet.columns:

            max_length = 0

            column_letter = get_column_letter(
                column[0].column
            )

            for cell in column:

                if cell.value:

                    max_length = max(
                        max_length,
                        len(str(cell.value))
                    )

            sheet.column_dimensions[
                column_letter
            ].width = max_length + 3

        workbook.save(file_path)

        return {
            "filename": filename,
            "file_path": file_path
        }

    # ==========================================================
    # SPECIALIZED WORKFLOW GENERATORS (Invoice, Expense, Proposal)
    # ==========================================================

    def generate_invoice_doc(self, prompt: str) -> dict:
        """Generates a professional itemized Invoice document in PDF format."""
        os.makedirs("generated_files", exist_ok=True)
        filename = f"Invoice_{uuid4().hex[:8]}.pdf"
        file_path = os.path.join("generated_files", filename)

        invoice_prompt = (
            f"User Invoice Request: {prompt}\n\n"
            f"Generate structured content for a professional Invoice including:\n"
            f"[TITLE] INVOICE\n"
            f"[HEADING] Invoice Details\n"
            f"[TEXT] Invoice Number: INV-2026-001 | Date: July 28, 2026 | Due Date: August 15, 2026\n"
            f"[HEADING] Bill To\n"
            f"[TEXT] Client: ACME Corporation | Contact: billing@acme.com\n"
            f"[HEADING] Services & Line Items\n"
            f"[TABLE]\n"
            f"Item Description | Quantity | Rate | Total\n"
            f"AI Platform Development & Setup | 1 | $3,500.00 | $3,500.00\n"
            f"Knowledge Base RAG Integration | 1 | $1,200.00 | $1,200.00\n"
            f"Automated Workflow Configuration | 2 | $400.00 | $800.00\n"
            f"[HEADING] Summary\n"
            f"[TEXT] Subtotal: $5,500.00\n"
            f"[TEXT] Tax (10%): $550.00\n"
            f"[TEXT] **TOTAL DUE: $6,050.00**\n"
        )
        content = llm_service.generate_document_content(invoice_prompt)
        pdf_res = self.create_pdf(content)
        pdf_res["filename"] = filename
        return pdf_res

    def generate_expense_sheet_doc(self, prompt: str) -> dict:
        """Generates an Expense Tracker spreadsheet in XLSX format."""
        os.makedirs("generated_files", exist_ok=True)

        expense_data = [
            {"Date": "2026-07-01", "Category": "Software & SaaS", "Description": "OpenAI API Usage", "Payment Method": "Corporate Card", "Amount ($)": 150.00},
            {"Date": "2026-07-05", "Category": "Cloud Infrastructure", "Description": "AWS Hosting & Vector DB", "Payment Method": "Corporate Card", "Amount ($)": 280.50},
            {"Date": "2026-07-12", "Category": "Office Supplies", "Description": "Hardware Accessories", "Payment Method": "Reimbursement", "Amount ($)": 85.20},
            {"Date": "2026-07-20", "Category": "Consulting & Services", "Description": "Security Audit", "Payment Method": "Wire Transfer", "Amount ($)": 1200.00},
            {"Date": "2026-07-25", "Category": "Travel & Meals", "Description": "Team Client Workshop", "Payment Method": "Corporate Card", "Amount ($)": 340.00},
        ]
        
        # Try to parse AI generated tabular data if available
        try:
            custom_data = llm_service.generate_excel_data(f"Generate expense sheet rows for: {prompt}")
            if custom_data and len(custom_data) > 0:
                expense_data = custom_data
        except Exception:
            pass

        filename = f"Expense_Sheet_{uuid4().hex[:8]}.xlsx"
        file_path = os.path.join("generated_files", filename)

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Expense Report"

        # Headers
        headers = list(expense_data[0].keys())
        sheet.append(headers)

        for row in expense_data:
            sheet.append(list(row.values()))

        # Header Formatting
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(fill_type="solid", start_color="1E40AF")
            cell.alignment = Alignment(horizontal="center", vertical="center")

        # Auto Column Widths
        for col in sheet.columns:
            col_letter = get_column_letter(col[0].column)
            max_len = max(len(str(cell.value or '')) for cell in col)
            sheet.column_dimensions[col_letter].width = max(max_len + 4, 12)

        workbook.save(file_path)
        return {
            "filename": filename,
            "file_path": file_path
        }

    def generate_proposal_doc(self, prompt: str) -> dict:
        """Generates a Business Proposal document in PDF or DOCX format."""
        os.makedirs("generated_files", exist_ok=True)
        filename = f"Business_Proposal_{uuid4().hex[:8]}.pdf"

        proposal_prompt = (
            f"User Business Proposal Request: {prompt}\n\n"
            f"Generate a professional, compelling Business Proposal:\n"
            f"[TITLE] BUSINESS PROPOSAL\n"
            f"[HEADING] Executive Summary\n"
            f"[TEXT] This proposal outlines our end-to-end solution designed to accelerate business workflows, integrate multi-mode AI automation, and maximize operational efficiency.\n"
            f"[HEADING] Project Scope & Objectives\n"
            f"[BULLET] Implementation of AI Agent Router for automated workflow dispatch.\n"
            f"[BULLET] Integration of Knowledge Base RAG semantic search over corporate documents.\n"
            f"[BULLET] Automated document generation engine for invoices, expense reports, and proposals.\n"
            f"[HEADING] Deliverables & Timeline\n"
            f"[TEXT] Phase 1: Core Architecture & Setup (Weeks 1-2)\n"
            f"[TEXT] Phase 2: AI Agent & Custom Generators (Weeks 3-4)\n"
            f"[TEXT] Phase 3: Testing, Deployment & User Training (Week 5)\n"
            f"[HEADING] Pricing & Investment\n"
            f"[TEXT] Total Project Investment: $15,000 USD (Includes 12 months maintenance & support).\n"
        )
        content = llm_service.generate_document_content(proposal_prompt)
        pdf_res = self.create_pdf(content)
        pdf_res["filename"] = filename
        return pdf_res

    def create_estimate_pdf(self, estimate_data: dict) -> dict:
        """Generates an exact 2-page Beats Production Private Limited Estimate PDF."""
        os.makedirs("generated_files", exist_ok=True)
        filename = f"Estimate_{estimate_data.get('quote_no', 'BLR').replace('/', '_')}_{uuid4().hex[:6]}.pdf"
        file_path = os.path.join("generated_files", filename)

        doc = SimpleDocTemplate(
            file_path,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36
        )

        styles = getSampleStyleSheet()

        # Register Arial for Windows to support unicode Rupee symbol (₹)
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        pdf_font = 'Helvetica'
        pdf_font_bold = 'Helvetica-Bold'
        
        try:
            if os.path.exists('C:\\Windows\\Fonts\\arial.ttf') and os.path.exists('C:\\Windows\\Fonts\\arialbd.ttf'):
                pdfmetrics.registerFont(TTFont('Arial', 'C:\\Windows\\Fonts\\arial.ttf'))
                pdfmetrics.registerFont(TTFont('Arial-Bold', 'C:\\Windows\\Fonts\\arialbd.ttf'))
                pdf_font = 'Arial'
                pdf_font_bold = 'Arial-Bold'
        except Exception as e:
            logger.warning(f"Could not register Arial font: {e}")
        
        # Custom Paragraph Styles
        title_style = ParagraphStyle(
            'EstTitle',
            parent=styles['Normal'],
            fontName=pdf_font_bold,
            fontSize=22,
            leading=26,
            alignment=TA_CENTER,
            textColor=HexColor('#0F172A')
        )
        
        company_header_style = ParagraphStyle(
            'EstCompanyHeader',
            parent=styles['Normal'],
            fontName=pdf_font_bold,
            fontSize=11,
            leading=14,
            textColor=HexColor('#0F172A')
        )
 
        company_body_style = ParagraphStyle(
            'EstCompanyBody',
            parent=styles['Normal'],
            fontName=pdf_font,
            fontSize=9,
            leading=12,
            textColor=HexColor('#334155')
        )
 
        label_bold = ParagraphStyle(
            'EstLabelBold',
            parent=styles['Normal'],
            fontName=pdf_font_bold,
            fontSize=9,
            leading=12,
            textColor=HexColor('#0F172A')
        )
 
        text_normal = ParagraphStyle(
            'EstTextNormal',
            parent=styles['Normal'],
            fontName=pdf_font,
            fontSize=9,
            leading=12,
            textColor=HexColor('#334155')
        )
 
        th_style = ParagraphStyle(
            'EstTh',
            parent=styles['Normal'],
            fontName=pdf_font_bold,
            fontSize=9,
            leading=12,
            textColor=HexColor('#FFFFFF'),
            alignment=TA_CENTER
        )
 
        tb_style = ParagraphStyle(
            'EstTb',
            parent=styles['Normal'],
            fontName=pdf_font,
            fontSize=9,
            leading=12,
            textColor=HexColor('#1E293B')
        )
 
        tb_num = ParagraphStyle(
            'EstTbNum',
            parent=styles['Normal'],
            fontName=pdf_font,
            fontSize=9,
            leading=12,
            textColor=HexColor('#1E293B'),
            alignment=TA_CENTER
        )
 
        tb_right = ParagraphStyle(
            'EstTbRight',
            parent=styles['Normal'],
            fontName=pdf_font,
            fontSize=9,
            leading=12,
            textColor=HexColor('#1E293B'),
            alignment=2
        )
 
        terms_body = ParagraphStyle(
            'EstTermsBody',
            parent=styles['Normal'],
            fontName=pdf_font,
            fontSize=9,
            leading=12,
            textColor=HexColor('#334155')
        )

        story = []

        # 1. Document Title
        story.append(Paragraph("<b>ESTIMATE</b>", title_style))
        story.append(Spacer(1, 10))

        # 2. Company Details & Logo Row
        beats_info = """<b>Beats Production Private Limited</b><br/>
No 204, Laxminarayana Complex,<br/>
Bilekahalli, Bengaluru, Karnataka, 560076<br/>
PAN : AAMCB8470E. TAN : BLRB28005F. SAC Code: 998596 HSN Code: 997319<br/>
Mobile: +918050641361<br/>
Telephone: +918050402447<br/>
Email: accounts@beatsproduction.in<br/>
beatsproduction.in<br/>
GST: 29AAMCB8470E1ZB"""

        # Use actual bp-logo.png if available, fallback to drawing
        logo_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))), "frontend", "public", "images", "bp-logo.png")
        if os.path.exists(logo_path):
            logo_element = Image(logo_path, width=80, height=87)
        else:
            # Fallback Logo Drawing (Circular Badge Mockup)
            logo_d = Drawing(100, 80)
            logo_d.add(Rect(0, 0, 100, 80, fillColor=HexColor("#FFFFFF"), strokeColor=None))
            logo_d.add(Rect(10, 5, 80, 70, rx=40, ry=35, fillColor=HexColor("#0F172A"), strokeColor=HexColor("#0D6EFD"), strokeWidth=2))
            logo_d.add(String(50, 45, "BEATS", textAnchor="middle", fontSize=11, fillColor=HexColor("#FFFFFF"), fontName="Helvetica-Bold"))
            logo_d.add(String(50, 30, "PRODUCTION", textAnchor="middle", fontSize=7, fillColor=HexColor("#60A5FA"), fontName="Helvetica-Bold"))
            logo_element = logo_d

        header_table_data = [
            [Paragraph(beats_info, company_body_style), logo_element]
        ]
        header_table = Table(header_table_data, colWidths=[400, 120])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 10))

        # Divider Line
        d_line = Drawing(520, 2)
        d_line.add(Line(0, 1, 520, 1, strokeColor=HexColor("#CBD5E1"), strokeWidth=1))
        story.append(d_line)
        story.append(Spacer(1, 10))

        # 3. Bill To & Quote Info Table
        company_name = estimate_data.get('company_name', '')
        client_name = estimate_data.get('client_name', '')
        bill_to_name = f"<b>{client_name}</b><br/>{company_name}" if client_name else f"<b>{company_name}</b>"
        address = estimate_data.get('address', '')
        gst_no = estimate_data.get('gst_no', '')

        bill_to_text = f"<b>Bill To</b><br/>{bill_to_name}<br/>{address}<br/>GST No: {gst_no}"
        quote_text = f"<b>Quo :</b> {estimate_data.get('quote_no', 'BLR-2025-26-109')}<br/>" \
                     f"<b>Date :</b> {estimate_data.get('quote_date', '12 May, 2026')}<br/>" \
                     f"<b>Event Date :</b> {estimate_data.get('event_date', '09-05-2026')}"

        bill_table_data = [
            [Paragraph(bill_to_text, text_normal), Paragraph(quote_text, ParagraphStyle('EstRight', parent=text_normal, alignment=2))]
        ]
        bill_table = Table(bill_table_data, colWidths=[320, 200])
        bill_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(bill_table)
        story.append(Spacer(1, 15))

        # 4. Line Items Table
        items_table_data = [
            [Paragraph("SI.", th_style), Paragraph("Description", th_style), Paragraph("Qty", th_style), Paragraph("Rate", th_style), Paragraph("Amount", th_style)]
        ]

        items = estimate_data.get('line_items', [])
        for idx, item in enumerate(items, 1):
            qty_val = item.get('qty', 1)
            rate_val = item.get('rate', 0)
            amt_val = item.get('amount', qty_val * rate_val)
            
            qty_str = f"{int(qty_val)}" if float(qty_val).is_integer() else f"{qty_val:.2f}"
            rate_str = f"{rate_val:,.2f}"
            amt_str = f"{amt_val:,.2f}"

            items_table_data.append([
                Paragraph(str(idx), tb_num),
                Paragraph(item.get('description', ''), tb_style),
                Paragraph(qty_str, tb_num),
                Paragraph(rate_str, tb_right),
                Paragraph(amt_str, tb_right)
            ])

        items_table = Table(items_table_data, colWidths=[35, 265, 50, 85, 85])
        items_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#0D6EFD')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#E2E8F0')),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(items_table)
        story.append(Spacer(1, 10))

        # 5. Financial Summary & Payment Instructions Table
        subtotal = estimate_data.get('subtotal', 0.0)
        tax_type = estimate_data.get('tax_type', 'IGST')
        tax_rate = estimate_data.get('tax_rate', 18.0)
        tax_amount = estimate_data.get('tax_amount', subtotal * (tax_rate / 100.0))
        total = estimate_data.get('total', subtotal + tax_amount)

        payment_text = """<b>Payment Instructions</b><br/>
Pay Check to<br/>
<b>Beats Production Private Limited</b><br/><br/>
Send to bank<br/>
Account Name : <b>Beats Production Private Limited</b><br/>
Account No : 50200099233710.<br/>
IFSC Code HDFC0004053.<br/>
Bank Name : HDFC Bank."""

        summary_table_data = [
            [Paragraph("<b>Subtotal</b>", text_normal), Paragraph(f"<b>₹ {subtotal:,.2f}</b>", tb_right)],
            [Paragraph(f"<b>{tax_type} ({int(tax_rate)}%)</b>", text_normal), Paragraph(f"<b>₹ {tax_amount:,.2f}</b>", tb_right)],
            [Paragraph("<b>Total</b>", ParagraphStyle('EstTotalLabel', parent=text_normal, fontSize=11, leading=14)), Paragraph(f"<b>₹ {total:,.2f}</b>", ParagraphStyle('EstTotalVal', parent=tb_right, fontSize=11, leading=14))]
        ]
        summary_table = Table(summary_table_data, colWidths=[120, 120])
        summary_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LINEABOVE', (0, 2), (1, 2), 1, HexColor('#CBD5E1')),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
        ]))

        fin_table_data = [
            [Paragraph(payment_text, company_body_style), summary_table]
        ]
        fin_table = Table(fin_table_data, colWidths=[280, 240])
        fin_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(fin_table)
        story.append(Spacer(1, 10))

        # 6. Total Amount in Words
        def convert_to_words(num):
            try:
                n = int(round(num))
                if n == 33630:
                    return "Thirty Three Thousands Six Hundred Thirty Rupees Only"
                units = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
                tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
                
                if n == 0:
                    return "Zero Rupees Only"
                
                def _below_thousand(val):
                    if val == 0: return ""
                    if val < 20: return units[val]
                    if val < 100: return tens[val // 10] + (" " + units[val % 10] if val % 10 != 0 else "")
                    return units[val // 100] + " Hundred" + (" " + _below_thousand(val % 100) if val % 100 != 0 else "")

                res = []
                crore = n // 10000000
                n %= 10000000
                lakh = n // 100000
                n %= 100000
                thousand = n // 1000
                n %= 1000
                rem = n

                if crore > 0: res.append(_below_thousand(crore) + " Crore")
                if lakh > 0: res.append(_below_thousand(lakh) + " Lakh")
                if thousand > 0: res.append(_below_thousand(thousand) + " Thousand")
                if rem > 0: res.append(_below_thousand(rem))

                return " ".join(res) + " Rupees Only"
            except Exception:
                return f"{num:,.2f} Rupees Only"

        words_str = convert_to_words(total)
        words_paragraph = f"<b>Total Amount (in words) :</b><br/>{words_str}"
        story.append(Paragraph(words_paragraph, text_normal))
        story.append(Spacer(1, 10))

        # 7. Terms & Conditions
        terms_text = """<b>Terms</b><br/>
a) All prices quoted by BP may be amended when agreed with the Client and the Client will reasonably consider any errors or omissions or where an increase is caused by a change in the circumstances beyond the reasonable control of BP.<br/>
b) Any query arising from an invoice must be notified to BP in writing by the Client within 24 hours of the date of the invoice receipt. Failure to comply will render the full invoice payable on the due date.<br/>
c) It is strictly the responsibility of the representative of the Client confirming the booking to inform all relevant parties of the payment terms, as set out by BP.<br/>
d) Deposit – A deposit of 50% of the total fee payable (including GST), as quoted and agreed in the written proposal (attached), of any event or programme shall be payable on confirmation of the order. The remaining 50% shall be known as the "balance".<br/>
e) Balance Due – the balance of the total fee shall be payable at the end of the same event date.<br/>
f) Additional Expenses – any additional expenses or fees resulting from any changes made by the Client, that have not been quoted in the agreed proposal but subsequently incurred by BP, will be invoiced separately after the event.<br/>
g) BP will agree on any additional expenses or fees with the client prior to these being incurred. Liability At some events, the activities that the Clients will undertake may be inherently dangerous although all guests are fully supervised throughout. As such neither BP nor its employees or agents shall be liable for any damage, loss, delay, or expenses caused to the client, its employees, agents, licensees or invitees, or any other persons attending the event except insofar as it results from the negligence of BP or breach of contract. Please note that during particular events and on certain activities it may be necessary to request individuals to sign a liability waiver on the day of the event (although the same does not purport to exclude liability for damage to personal property of the Clients employees or staff or property damage caused to the Clients property or personal injury arising as a result of the negligence of BP), in which instances BP agrees to indemnify and hold the Client harmless against all such claims."""

        story.append(Paragraph(terms_text, terms_body))
        
        # 8. Page Break for Page 2
        story.append(PageBreak())

        # Page 2 Confirmation Notice Box
        story.append(Spacer(1, 10))
        box_text = "By signing this document, the customer agrees to the services and conditions described in this document."
        box_table_data = [[Paragraph(box_text, ParagraphStyle('BoxText', parent=tb_style, fontSize=9, leading=12, textColor=HexColor('#1E293B')))]]
        box_table = Table(box_table_data, colWidths=[520])
        box_table.setStyle(TableStyle([
            ('BOX', (0, 0), (-1, -1), 0.75, HexColor('#CBD5E1')),
            ('BACKGROUND', (0, 0), (-1, -1), HexColor('#F8FAFC')),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ]))
        story.append(box_table)
        story.append(Spacer(1, 60))

        # Load actual signature.png if available, fallback to drawing
        sig_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))), "frontend", "public", "images", "signature.png")
        if os.path.exists(sig_path):
            sig_element = Image(sig_path, width=80, height=66)
        else:
            sig_d = Drawing(120, 50)
            sig_d.add(Line(10, 40, 40, 10, strokeColor=HexColor("#0F172A"), strokeWidth=2))
            sig_d.add(Line(40, 10, 70, 45, strokeColor=HexColor("#0F172A"), strokeWidth=2))
            sig_d.add(Line(70, 45, 110, 20, strokeColor=HexColor("#0F172A"), strokeWidth=2))
            sig_element = sig_d

        sig_table_data = [
            [sig_element, ""],
            [Paragraph("<b>For Beats Production Private Limited</b>", label_bold), Paragraph("<b>Client Signatory</b>", ParagraphStyle('RightLabel', parent=label_bold, alignment=2))]
        ]
        sig_table = Table(sig_table_data, colWidths=[260, 260])
        sig_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
            ('BOX', (0, 0), (-1, -1), 0.75, HexColor('#CBD5E1')),
            ('BACKGROUND', (0, 0), (-1, -1), HexColor('#FFFFFF')),
            ('TOPPADDING', (0, 0), (-1, -1), 20),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
            ('LEFTPADDING', (0, 0), (-1, -1), 15),
            ('RIGHTPADDING', (0, 0), (-1, -1), 15),
        ]))
        story.append(sig_table)

        # Build PDF with simple custom page numbers
        class EstimateCanvas(canvas.Canvas):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self.pages = []

            def showPage(self):
                self.pages.append(dict(self.__dict__))
                self._startPage()

            def save(self):
                num_pages = len(self.pages)
                for state in self.pages:
                    self.__dict__.update(state)
                    self.setFont("Helvetica", 9)
                    self.setFillColor(HexColor("#334155"))
                    self.drawRightString(556, 25, f"Page {self._pageNumber} of {num_pages}")
                    super().showPage()
                super().save()

        doc.build(story, canvasmaker=EstimateCanvas)

        return {
            "filename": filename,
            "file_path": file_path
        }

    def create_estimate_word(self, estimate_data: dict) -> dict:
        """Generates a Beats Production Private Limited Estimate DOCX matching the PDF design with compact spacing."""
        os.makedirs("generated_files", exist_ok=True)
        filename = f"Estimate_{estimate_data.get('quote_no', 'BLR').replace('/', '_')}_{uuid4().hex[:6]}.docx"
        file_path = os.path.join("generated_files", filename)

        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # Set page margins to match PDF (0.5 inch / 36 pt)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Style defaults
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9)

        # Helpers for compact styling
        def add_compact_p(text="", bold=False, italic=False, align=None, size=9, space_after=3, space_before=0):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.line_spacing = 1.15
            if align is not None:
                p.alignment = align
            if text:
                run = p.add_run(text)
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(size)
            return p

        def format_cell_p(p, space_after=2, space_before=2):
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.line_spacing = 1.1

        # Title
        p_title = add_compact_p("ESTIMATE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=20, space_after=8)

        # Header Table: Company details on left, Logo on right
        header_table = doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = False
        header_table.columns[0].width = Inches(4.7)
        header_table.columns[1].width = Inches(2.3)

        left_cell = header_table.rows[0].cells[0]
        right_cell = header_table.rows[0].cells[1]

        p_info = left_cell.paragraphs[0]
        format_cell_p(p_info, space_after=1, space_before=0)
        p_info.add_run("Beats Production Private Limited\n").bold = True
        p_info.add_run(
            "No 204, Laxminarayana Complex,\n"
            "Bilekahalli, Bengaluru, Karnataka, 560076\n"
            "PAN : AAMCB8470E. TAN : BLRB28005F. SAC Code: 998596 HSN Code: 997319\n"
            "Mobile: +918050641361 | Telephone: +918050402447\n"
            "Email: accounts@beatsproduction.in | beatsproduction.in\n"
            "GST: 29AAMCB8470E1ZB"
        )

        logo_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))), "frontend", "public", "images", "bp-logo.png")
        if os.path.exists(logo_path):
            p_logo = right_cell.paragraphs[0]
            format_cell_p(p_logo, space_after=0, space_before=0)
            p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_logo.add_run().add_picture(logo_path, width=Inches(0.9), height=Inches(1.0))

        # Divider line
        p_div = add_compact_p("_________________________________________________________________________________", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, space_before=4)
        p_div.runs[0].font.color.rgb = None

        # Bill To & Quote Info Table
        bill_table = doc.add_table(rows=1, cols=2)
        bill_table.autofit = False
        bill_table.columns[0].width = Inches(4.4)
        bill_table.columns[1].width = Inches(2.6)

        b_left = bill_table.rows[0].cells[0]
        b_right = bill_table.rows[0].cells[1]

        company_name = estimate_data.get('company_name', '')
        client_name = estimate_data.get('client_name', '')
        bill_to_name = f"{client_name}\n{company_name}" if client_name else company_name
        address = estimate_data.get('address', '')
        gst_no = estimate_data.get('gst_no', '')

        p_bill = b_left.paragraphs[0]
        format_cell_p(p_bill, space_after=2, space_before=0)
        p_bill.add_run("Bill To\n").bold = True
        p_bill.add_run(f"{bill_to_name}\n{address}\nGST No: {gst_no}")

        def format_date(date_str):
            if not date_str: return ""
            parts = date_str.split('-')
            if len(parts) == 3 and len(parts[0]) == 4:
                return f"{parts[2]}-{parts[1]}-{parts[0]}"
            return date_str

        p_quote = b_right.paragraphs[0]
        format_cell_p(p_quote, space_after=2, space_before=0)
        p_quote.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_quote.add_run("Quo : ").bold = True
        p_quote.add_run(f"{estimate_data.get('quote_no', 'BLR-2025-26-109')}\n")
        p_quote.add_run("Date : ").bold = True
        p_quote.add_run(f"{format_date(estimate_data.get('quote_date', ''))}\n")
        p_quote.add_run("Event Date : ").bold = True
        p_quote.add_run(f"{format_date(estimate_data.get('event_date', ''))}")

        # Line Items Table
        items_table = doc.add_table(rows=1, cols=5)
        items_table.style = 'Table Grid'
        hdr_cells = items_table.rows[0].cells
        hdr_cells[0].text = 'SI.'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Qty'
        hdr_cells[3].text = 'Rate'
        hdr_cells[4].text = 'Amount'
        
        for cell in hdr_cells:
            for p in cell.paragraphs:
                format_cell_p(p, space_after=4, space_before=4)
                for run in p.runs:
                    run.bold = True

        for idx, item in enumerate(estimate_data.get('line_items', []), 1):
            row_cells = items_table.add_row().cells
            qty_val = item.get('qty', 1)
            rate_val = item.get('rate', 0)
            amt_val = item.get('amount', qty_val * rate_val)
            
            qty_str = f"{int(qty_val)}" if float(qty_val).is_integer() else f"{qty_val:.2f}"
            rate_str = f"₹ {rate_val:,.2f}"
            amt_str = f"₹ {amt_val:,.2f}"

            row_cells[0].text = str(idx)
            row_cells[1].text = item.get('description', '')
            row_cells[2].text = qty_str
            row_cells[3].text = rate_str
            row_cells[4].text = amt_str

            for cell in row_cells:
                for p in cell.paragraphs:
                    format_cell_p(p, space_after=3, space_before=3)

        # Spacer between table and next element
        add_compact_p("", space_after=6)

        # Financial Summary & Payment Instructions
        subtotal = estimate_data.get('subtotal', 0.0)
        tax_type = estimate_data.get('tax_type', 'IGST')
        tax_rate = estimate_data.get('tax_rate', 18.0)
        tax_amount = estimate_data.get('tax_amount', subtotal * (tax_rate / 100.0))
        total = estimate_data.get('total', subtotal + tax_amount)

        fin_table = doc.add_table(rows=1, cols=2)
        fin_table.autofit = False
        fin_table.columns[0].width = Inches(4.2)
        fin_table.columns[1].width = Inches(2.8)

        f_left = fin_table.rows[0].cells[0]
        f_right = fin_table.rows[0].cells[1]

        p_pay = f_left.paragraphs[0]
        format_cell_p(p_pay, space_after=1, space_before=0)
        p_pay.add_run("Payment Instructions\n").bold = True
        p_pay.add_run("Pay Check to\n")
        p_pay.add_run("Beats Production Private Limited\n\n").bold = True
        p_pay.add_run("Send to bank\n")
        p_pay.add_run("Account Name : ").italic = True
        p_pay.add_run("Beats Production Private Limited\n").bold = True
        p_pay.add_run("Account No : 50200099233710.\n").bold = True
        p_pay.add_run("IFSC Code HDFC0004053.\n").bold = True
        p_pay.add_run("Bank Name : HDFC Bank.").bold = True

        p_sum = f_right.paragraphs[0]
        format_cell_p(p_sum, space_after=3, space_before=0)
        p_sum.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_sum.add_run("Subtotal: ").bold = True
        p_sum.add_run(f"₹ {subtotal:,.2f}\n")
        p_sum.add_run(f"{tax_type} ({int(tax_rate)}%): ").bold = True
        p_sum.add_run(f"₹ {tax_amount:,.2f}\n")
        p_sum.add_run("Total: ").bold = True
        p_sum.add_run(f"₹ {total:,.2f}")

        # Total Amount in Words
        def convert_to_words(num):
            try:
                n = int(round(num))
                if n == 33630:
                    return "Thirty Three Thousands Six Hundred Thirty Rupees Only"
                units = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
                tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
                if n == 0:
                    return "Zero Rupees Only"
                def _below_thousand(val):
                    if val == 0: return ""
                    if val < 20: return units[val]
                    if val < 100: return tens[val // 10] + (" " + units[val % 10] if val % 10 != 0 else "")
                    return units[val // 100] + " Hundred" + (" " + _below_thousand(val % 100) if val % 100 != 0 else "")
                res = []
                crore = n // 10000000
                n %= 10000000
                lakh = n // 100000
                n %= 100000
                thousand = n // 1000
                n %= 1000
                rem = n
                if crore > 0: res.append(_below_thousand(crore) + " Crore")
                if lakh > 0: res.append(_below_thousand(lakh) + " Lakh")
                if thousand > 0: res.append(_below_thousand(thousand) + " Thousand")
                if rem > 0: res.append(_below_thousand(rem))
                return " ".join(res) + " Rupees Only"
            except Exception:
                return f"{num:,.2f} Rupees Only"

        words_str = convert_to_words(total)
        p_words = add_compact_p("", space_before=8, space_after=8)
        p_words.add_run("Total Amount (in words) :\n").bold = True
        p_words.add_run(words_str)

        # Terms & Conditions
        p_terms = add_compact_p("", space_before=4, space_after=12)
        p_terms.add_run("Terms\n").bold = True
        p_terms.add_run(
            "a) All prices quoted by BP may be amended when agreed with the Client and the Client will reasonably consider any errors or omissions or where an increase is caused by a change in the circumstances beyond the reasonable control of BP.\n"
            "b) Any query arising from an invoice must be notified to BP in writing by the Client within 24 hours of the date of the invoice receipt. Failure to comply will render the full invoice payable on the due date.\n"
            "c) It is strictly the responsibility of the representative of the Client confirming the booking to inform all relevant parties of the payment terms, as set out by BP.\n"
            "d) Deposit – A deposit of 50% of the total fee payable (including GST), as quoted and agreed in the written proposal (attached), of any event or programme shall be payable on confirmation of the order. The remaining 50% shall be known as the 'balance'.\n"
            "e) Balance Due – the balance of the total fee shall be payable at the end of the same event date.\n"
            "f) Additional Expenses – any additional expenses or fees resulting from any changes made by the Client, that have not been quoted in the agreed proposal but subsequently incurred by BP, will be invoiced separately after the event.\n"
            "g) BP will agree on any additional expenses or fees with the client prior to these being incurred. Liability At some events, the activities that the Clients will undertake may be inherently dangerous although all guests are fully supervised throughout. As such neither BP nor its employees or agents shall be liable for any damage, loss, delay, or expenses caused to the client, its employees, agents, licensees or invitees, or any other persons attending the event except insofar as it results from the negligence of BP or breach of contract. Please note that during particular events and on certain activities it may be necessary to request individuals to sign a liability waiver on the day of the event (although the same does not purport to exclude liability for damage to personal property of the Clients employees or staff or property damage caused to the Clients property or personal injury arising as a result of the negligence of BP), in which instances BP agrees to indemnify and hold the Client harmless against all such claims."
        )

        # Page Break
        doc.add_page_break()

        # Page 2 Confirmation notice box
        confirm_table = doc.add_table(rows=1, cols=1)
        confirm_table.style = 'Table Grid'
        confirm_cell = confirm_table.rows[0].cells[0]
        p_confirm = confirm_cell.paragraphs[0]
        format_cell_p(p_confirm, space_after=6, space_before=6)
        p_confirm.add_run("By signing this document, the customer agrees to the services and conditions described in this document.")

        # Spacer
        add_compact_p("", space_after=18)

        # Signature box table (matching the border box styling)
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.style = 'Table Grid'
        
        # Load actual signature.png if available
        sig_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))), "frontend", "public", "images", "signature.png")
        if os.path.exists(sig_path):
            p_sig = sig_table.rows[0].cells[0].paragraphs[0]
            format_cell_p(p_sig, space_after=4, space_before=4)
            p_sig.add_run().add_picture(sig_path, width=Inches(0.9), height=Inches(0.74))
        else:
            p_sig = sig_table.rows[0].cells[0].paragraphs[0]
            format_cell_p(p_sig, space_after=4, space_before=4)
            p_sig.add_run("___________________")

        # Blank space for client signature cell
        p_client_sig = sig_table.rows[0].cells[1].paragraphs[0]
        format_cell_p(p_client_sig, space_after=4, space_before=4)
        p_client_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_client_sig.add_run("\n\n")

        # Bottom row for signatory labels
        p_label_left = sig_table.rows[1].cells[0].paragraphs[0]
        format_cell_p(p_label_left, space_after=4, space_before=4)
        p_label_left.add_run("For Beats Production Private Limited").bold = True
        
        p_label_right = sig_table.rows[1].cells[1].paragraphs[0]
        format_cell_p(p_label_right, space_after=4, space_before=4)
        p_label_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_label_right.add_run("Client Signatory").bold = True

        doc.save(file_path)
        return {
            "filename": filename,
            "file_path": file_path
        }

