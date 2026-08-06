import os
from docx import Document
from docx.shared import Pt, Inches

def create_onboarding_doc(output_path="Onboarding_Guide.docx"):
    doc = Document()

    # Title
    title = doc.add_heading('AI Workspace - Project Onboarding Guide', 0)
    title.alignment = 1 # Center

    # Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Welcome to the AI Workspace project! This guide will help you understand the core architecture, "
        "technologies, and file structure of our application. AI Workspace is a modern full-stack application "
        "designed to provide AI-powered features like chat, knowledge assistance, document analysis, and estimate generation."
    )

    # Architecture Overview
    doc.add_heading('2. Architecture Overview', level=1)
    doc.add_paragraph("The project is divided into two main components: a React-based frontend and a Python FastAPI backend.")
    
    doc.add_heading('Frontend Technology Stack:', level=2)
    ul1 = doc.add_paragraph(style='List Bullet')
    ul1.add_run("React & Vite: ").bold = True
    ul1.add_run("For fast development and optimized production builds.")
    ul2 = doc.add_paragraph(style='List Bullet')
    ul2.add_run("Material UI (MUI): ").bold = True
    ul2.add_run("Provides the core UI components and styling (theme-based).")
    ul3 = doc.add_paragraph(style='List Bullet')
    ul3.add_run("React Router: ").bold = True
    ul3.add_run("Handles client-side routing and protected routes.")
    ul4 = doc.add_paragraph(style='List Bullet')
    ul4.add_run("Axios: ").bold = True
    ul4.add_run("For making HTTP requests to the backend API.")

    doc.add_heading('Backend Technology Stack:', level=2)
    ul5 = doc.add_paragraph(style='List Bullet')
    ul5.add_run("FastAPI: ").bold = True
    ul5.add_run("A modern, fast web framework for building APIs with Python.")
    ul6 = doc.add_paragraph(style='List Bullet')
    ul6.add_run("SQLAlchemy & PostgreSQL/SQLite: ").bold = True
    ul6.add_run("ORM and database layer for data persistence.")
    ul7 = doc.add_paragraph(style='List Bullet')
    ul7.add_run("AI/LLM Integrations: ").bold = True
    ul7.add_run("Uses OpenAI, Tavily (for web search), and document parsing libraries (pymupdf, python-docx, openpyxl).")

    # Directory Structure
    doc.add_heading('3. Directory & File Structure', level=1)
    doc.add_paragraph("The workspace is located at c:\\projects\\AI-Workspace and is split into 'frontend' and 'backend' directories.")

    doc.add_heading('Frontend (frontend/)', level=2)
    p_fe = doc.add_paragraph()
    p_fe.add_run("src/App.jsx & main.jsx: ").bold = True
    p_fe.add_run("The entry points where the app is initialized, theme is applied, and routes are wrapped.\n")
    p_fe.add_run("src/routes/AppRoutes.jsx: ").bold = True
    p_fe.add_run("Defines all the routes like /chat, /files, /knowledge, /estimate. It uses a ProtectedRoute component for authenticated areas.\n")
    p_fe.add_run("src/pages/: ").bold = True
    p_fe.add_run("Contains top-level components representing different views (e.g., Dashboard, Chat, Settings).\n")
    p_fe.add_run("src/components/: ").bold = True
    p_fe.add_run("Reusable UI components and layouts (like MainLayout).\n")
    p_fe.add_run("src/services/: ").bold = True
    p_fe.add_run("API client functions using Axios to communicate with the backend.")

    doc.add_heading('Backend (backend/app/)', level=2)
    p_be = doc.add_paragraph()
    p_be.add_run("main.py: ").bold = True
    p_be.add_run("The FastAPI application instance, CORS configuration, and route inclusions.\n")
    p_be.add_run("api/: ").bold = True
    p_be.add_run("Contains routers (auth.py, chat.py, file.py, knowledge.py, estimate.py) that define the API endpoints.\n")
    p_be.add_run("models/: ").bold = True
    p_be.add_run("SQLAlchemy database models (user.py, chat.py, message.py, document.py).\n")
    p_be.add_run("services/: ").bold = True
    p_be.add_run("The core business logic resides here (e.g., llm_service.py, file_service.py, rag_service.py). Controllers in api/ call these services.\n")
    p_be.add_run("database/: ").bold = True
    p_be.add_run("Database connection and session management.\n")
    p_be.add_run("core/: ").bold = True
    p_be.add_run("Configuration settings and environment variables.")

    # How it connects
    doc.add_heading('4. How the Flow Works (Data Connection)', level=1)
    flow_text = (
        "1. User Action: A user interacts with a frontend component (e.g., clicks 'Send' in the Chat interface).\n"
        "2. API Call: The React component calls a function in frontend/src/services/, which makes an Axios HTTP request to the backend.\n"
        "3. Routing: The request hits FastAPI (backend/app/main.py) and is routed to the appropriate endpoint in backend/app/api/.\n"
        "4. Business Logic & DB: The API endpoint extracts the data and calls a function in backend/app/services/. This service interacts with the database via SQLAlchemy models (backend/app/models/) and external APIs (like OpenAI).\n"
        "5. Response: The service returns the processed data (or an AI-generated response) back to the router, which sends a JSON response to the frontend.\n"
        "6. UI Update: The React component updates its state and re-renders the UI with the new data."
    )
    doc.add_paragraph(flow_text)

    # Getting Started
    doc.add_heading('5. Getting Started (Running Locally)', level=1)
    
    doc.add_heading('Backend:', level=2)
    b_steps = doc.add_paragraph()
    b_steps.add_run("1. Navigate to the backend directory.\n")
    b_steps.add_run("2. Create and activate a virtual environment.\n")
    b_steps.add_run("3. Run: ")
    b_steps.add_run("pip install -r requirements.txt\n").italic = True
    b_steps.add_run("4. Run: ")
    b_steps.add_run("uvicorn app.main:app --reload\n").italic = True
    
    doc.add_heading('Frontend:', level=2)
    f_steps = doc.add_paragraph()
    f_steps.add_run("1. Navigate to the frontend directory.\n")
    f_steps.add_run("2. Run: ")
    f_steps.add_run("npm install\n").italic = True
    f_steps.add_run("3. Run: ")
    f_steps.add_run("npm run dev\n").italic = True

    # Save the document
    try:
        doc.save(output_path)
        print(f"Document successfully generated at {output_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    import os
    workspace_dir = r"c:\projects\AI-Workspace"
    output_file = os.path.join(workspace_dir, "Project_Onboarding_Guide.docx")
    create_onboarding_doc(output_file)
